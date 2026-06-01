"""
docx_reader.py
─────────────────────────────────────────────────────────────────────────────
Extracts raw text from DOCX resumes.

Extraction order (all layers joined in document reading order)
--------------------------------------------------------------
1. BODY PARAGRAPHS  — Normal content: headings, bullets, plain text.
2. TABLES           — Each cell is read left-to-right, top-to-bottom.
                      Table text is interleaved at the position it appears
                      in the XML, not appended at the end.
3. HEADERS          — Extracted from every section's header part.
4. FOOTERS          — Extracted from every section's footer part.
5. TEXT BOXES       — Inline shapes / drawing canvas paragraphs (sdtContent
                      + txbxContent), common in designed resume templates.

Why all layers?
    Resumes often place the candidate's name, contact info, or sections
    inside text boxes or table cells. Skipping any layer means missing
    data the downstream entity extractor needs.

Public API
----------
    extract_text_from_docx(file_path: str | Path) -> DOCXResult

    DOCXResult  (TypedDict)
    ├── success      : bool
    ├── text         : str          # full document text, "" on failure
    ├── paragraphs   : int          # body paragraph count (0 on failure)
    ├── tables       : int          # table count (0 on failure)
    └── error        : str | None   # human-readable message on failure
"""

from __future__ import annotations

import logging
from pathlib import Path
from typing import TypedDict
from xml.etree import ElementTree as ET

from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from docx.table import Table
from docx.text.paragraph import Paragraph

# ── Logging ──────────────────────────────────────────────────────────────────
logger = logging.getLogger(__name__)
logger.addHandler(logging.NullHandler())

# Word Processing ML namespace — required for raw XML traversal
_WNS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


# ── Return type ──────────────────────────────────────────────────────────────
class DOCXResult(TypedDict):
    success: bool
    text: str
    paragraphs: int
    tables: int
    error: str | None


# ── Internal helpers ─────────────────────────────────────────────────────────

def _validate_path(file_path: str | Path) -> tuple[Path, DOCXResult | None]:
    path = Path(file_path).resolve()

    if not path.exists():
        return path, _error_result(f"File not found: {path}")
    if not path.is_file():
        return path, _error_result(f"Path is not a file: {path}")
    if path.suffix.lower() not in (".docx", ".docm"):
        return path, _error_result(
            f"Unsupported extension '{path.suffix}'. Expected .docx or .docm"
        )
    if path.stat().st_size == 0:
        return path, _error_result(f"File is empty: {path}")

    return path, None


def _error_result(message: str) -> DOCXResult:
    logger.error(message)
    return DOCXResult(
        success=False,
        text="",
        paragraphs=0,
        tables=0,
        error=message,
    )


def _para_text(para: Paragraph) -> str:
    """
    Concatenate all runs in a paragraph.

    Uses the paragraph's internal XML to also capture hyperlink text
    (stored inside <w:hyperlink> child elements which python-docx's
    .text property misses in some versions).
    """
    try:
        return "".join(
            node.text or ""
            for node in para._p.iter(f"{_WNS}t")
        )
    except Exception:  # noqa: BLE001
        return para.text or ""


def _table_text(table: Table) -> list[str]:
    """
    Extract text from every cell, row-by-row.
    Returns a flat list of non-empty cell strings.
    """
    cell_texts: list[str] = []
    try:
        for row in table.rows:
            for cell in row.cells:
                parts = [_para_text(p) for p in cell.paragraphs]
                cell_str = " ".join(p for p in parts if p.strip())
                if cell_str.strip():
                    cell_texts.append(cell_str.strip())
    except Exception as exc:  # noqa: BLE001
        logger.warning("Table cell extraction error (partial results kept): %s", exc)
    return cell_texts


def _header_footer_text(doc: Document) -> list[str]:
    """
    Extract text from all section headers and footers.
    python-docx exposes these via section.header / section.footer.
    """
    texts: list[str] = []
    try:
        for section in doc.sections:
            for hf in (section.header, section.footer,
                       section.even_page_header, section.even_page_footer,
                       section.first_page_header, section.first_page_footer):
                try:
                    if hf is not None and hf.is_linked_to_previous is False:
                        for para in hf.paragraphs:
                            t = _para_text(para)
                            if t.strip():
                                texts.append(t.strip())
                except Exception as exc:  # noqa: BLE001
                    logger.debug("Header/footer section error: %s", exc)
    except Exception as exc:  # noqa: BLE001
        logger.warning("Header/footer extraction failed: %s", exc)
    return texts


def _text_box_text(doc: Document) -> list[str]:
    """
    Extract text from drawing text boxes (txbxContent) and
    structured document tags (sdtContent).

    Many designer resume templates place name/contact info inside
    floating text boxes. python-docx's normal API skips these entirely.
    We walk the raw XML of the body element to find them.
    """
    texts: list[str] = []
    _TXBX  = f"{_WNS}txbxContent"
    _SDT   = f"{_WNS}sdtContent"
    _T     = f"{_WNS}t"

    try:
        body_xml = doc.element.body
        for container_tag in (_TXBX, _SDT):
            for container in body_xml.iter(container_tag):
                for t_node in container.iter(_T):
                    val = (t_node.text or "").strip()
                    if val:
                        texts.append(val)
    except Exception as exc:  # noqa: BLE001
        logger.warning("Text box extraction failed: %s", exc)
    return texts


# ── Body block iterator (preserves document order) ───────────────────────────

def _iter_body_blocks(doc: Document):
    """
    Yield body-level blocks in document XML order.

    python-docx exposes doc.paragraphs and doc.tables separately,
    losing the interleaving. This reconstructs the original order by
    walking the body element directly.
    """
    body = doc.element.body
    para_tag  = f"{_WNS}p"
    table_tag = f"{_WNS}tbl"

    for child in body:
        tag = child.tag
        if tag == para_tag:
            yield Paragraph(child, doc)
        elif tag == table_tag:
            yield Table(child, doc)


# ── Public API ────────────────────────────────────────────────────────────────

def extract_text_from_docx(file_path: str | Path) -> DOCXResult:
    """
    Extract all text from a DOCX file.

    Covers: body paragraphs, table cells, headers, footers, and text boxes.

    Parameters
    ----------
    file_path : str | Path
        Path to a .docx or .docm file.

    Returns
    -------
    DOCXResult
        - success    (bool)    : True if any text was extracted.
        - text       (str)     : Full document text joined by newlines.
        - paragraphs (int)     : Number of body-level paragraphs.
        - tables     (int)     : Number of body-level tables.
        - error      (str|None): Error description when success=False.

    Notes
    -----
    - Legacy .doc files are NOT supported; convert them to .docx first.
    - Password-protected files return success=False.
    - Empty documents (no text anywhere) return success=False.
    """
    path, validation_error = _validate_path(file_path)
    if validation_error is not None:
        return validation_error

    logger.info("Starting DOCX extraction: %s", path.name)

    # ── Open document ─────────────────────────────────────────────────────────
    try:
        doc = Document(str(path))
    except PackageNotFoundError:
        msg = (
            f"'{path.name}' could not be opened as a DOCX package. "
            "It may be a legacy .doc file, corrupt, or password-protected."
        )
        return _error_result(msg)
    except Exception as exc:  # noqa: BLE001
        return _error_result(
            f"Failed to open '{path.name}': {exc}"
        )

    # ── Body: paragraphs + tables in document order ───────────────────────────
    body_lines: list[str] = []
    para_count = 0
    table_count = 0

    try:
        for block in _iter_body_blocks(doc):
            if isinstance(block, Paragraph):
                para_count += 1
                t = _para_text(block)
                if t.strip():
                    body_lines.append(t.strip())
            elif isinstance(block, Table):
                table_count += 1
                cell_texts = _table_text(block)
                body_lines.extend(cell_texts)
    except Exception as exc:  # noqa: BLE001
        logger.warning("Body block iteration error (partial results kept): %s", exc)

    # ── Headers + footers ─────────────────────────────────────────────────────
    hf_lines = _header_footer_text(doc)

    # ── Text boxes ────────────────────────────────────────────────────────────
    txbx_lines = _text_box_text(doc)

    # ── Assemble final text ───────────────────────────────────────────────────
    # Prepend header/footer + text-box content so contact info found there
    # appears near the top (where the name extractor expects it).
    all_lines = hf_lines + txbx_lines + body_lines
    full_text = "\n".join(line for line in all_lines if line.strip())

    if not full_text.strip():
        msg = (
            f"No text found in '{path.name}'. "
            "The document may be empty or contain only images/objects."
        )
        return DOCXResult(
            success=False,
            text="",
            paragraphs=para_count,
            tables=table_count,
            error=msg,
        )

    logger.info(
        "DOCX extraction succeeded: %d paragraphs, %d tables, %d chars.",
        para_count, table_count, len(full_text),
    )
    return DOCXResult(
        success=True,
        text=full_text,
        paragraphs=para_count,
        tables=table_count,
        error=None,
    )